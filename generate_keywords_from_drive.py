import os
import logging
import subprocess
import sys

# 🔒 Автоустановка PyPDF2 и python-docx, если не установлены
try:
    from PyPDF2 import PdfReader
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "PyPDF2"])
    from PyPDF2 import PdfReader

try:
    from docx import Document
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

from config import PROGRAM_SHEETS, PROGRAM_SHEETS_LIST, OPENAI_API_KEY
from services.google_sheets_service import get_sheet_data, update_keywords_for_discipline
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from openai import OpenAI
import mimetypes
import io
import re

# ID папок на Google Диске по программам
DRIVE_FOLDER_IDS = {
    "МРК": "1QXr2qpizqjcn3O3F8_PHu-_b-5Es8Dds",
    "ТПР": "1fO3VSIsE4C1B2n2EAlKmOM_g9Qgiqjum",
    "БХ": "1HvR5wqnPqupqKtTqG8yOxzBJnMFO0YQY"
}

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Авторизация Google Drive
SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]
SERVICE_ACCOUNT_FILE = "credentials.json"
creds = service_account.Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)
drive_service = build("drive", "v3", credentials=creds)

# OpenAI
client = OpenAI(api_key=OPENAI_API_KEY)

def extract_text_from_docx(file_bytes, file_name):
    file_stream = io.BytesIO(file_bytes)
    doc = Document(file_stream)

    full_text = "\n".join([para.text for para in doc.paragraphs])
    content_blocks = []

    # === РПД: Извлекаем 2-й столбец всех таблиц ===
    if file_name.lower().startswith("рпд"):
        for table in doc.tables:
            for row in table.rows[1:]:
                if len(row.cells) > 1:
                    cell = row.cells[1]
                    text = cell.text.strip()
                    if text:
                        content_blocks.append(text)
        combined_text = "\n".join(content_blocks)

    # === РПП ===
    elif file_name.lower().startswith("рпп"):
        pattern = r"(?<=структура и содержание).*?(?=8\\.|8 |формы отчетности|оформление документации)"
        match = re.search(pattern, full_text, re.IGNORECASE | re.DOTALL)
        combined_text = match.group(0).strip() if match else full_text

    # === ГИА ===
    elif file_name.lower().startswith("гиа"):
        pattern = r"(?<=требования к вкр).*?(?=4\\.|перечень тем|5\\.|критерии оценки)"
        match = re.search(pattern, full_text, re.IGNORECASE | re.DOTALL)
        combined_text = match.group(0).strip() if match else full_text

    else:
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    table_text.append(" | ".join(row_text))
        combined_text = full_text + "\n" + "\n".join(table_text)

    return combined_text

def generate_keywords_from_text(text, max_words=200):
    prompt = (
        f"Ты — методист. Прочитай фрагмент дисциплины и выдели {max_words} ключевых слов (через запятую), "
        f"характеризующих её содержание. Не добавляй пояснений.\n\n{text[:3000]}"
    )

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Ты — методист, специалист по учебным дисциплинам."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logger.error(f"Ошибка OpenAI: {e}")
        return ""

def find_drive_file(folder_id, discipline_name):
    simplified = discipline_name.replace(" ", "_").lower()
    query = f"'{folder_id}' in parents and trashed = false"
    results = drive_service.files().list(q=query, fields="files(id, name)").execute()
    for file in results.get("files", []):
        if simplified in file['name'].lower():
            return file
    return None

def download_file_content(file_id):
    request = drive_service.files().get_media(fileId=file_id)
    file_data = io.BytesIO()
    downloader = MediaIoBaseDownload(file_data, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()
    return file_data.getvalue()

def process_program(program_name, sheet_name):
    logger.info(f"📘 Обрабатываем программу: {program_name}")
    folder_id = DRIVE_FOLDER_IDS.get(program_name)
    not_updated = []

    data = get_sheet_data(PROGRAM_SHEETS, f"{sheet_name}!A2:C")
    for i, row in enumerate(data, start=2):
        if len(row) < 2:
            continue
        current_keywords = row[2] if len(row) >= 3 else ""
        if current_keywords.strip():
            logger.info(f"⏩ Пропуск: ключевые слова уже есть для {row[1]}")
            continue
        module, discipline = row[0], row[1]
        logger.info(f"🔍 {module} → {discipline}")

        file = find_drive_file(folder_id, discipline)
        if not file:
            logger.warning(f"⚠️ Не найден файл для: {discipline}")
            not_updated.append((module, discipline))
            continue

        content = download_file_content(file['id'])
        filename = file['name'].lower()

        if filename.endswith(".docx"):
            text = extract_text_from_docx(content, filename)
        else:
            logger.warning(f"⚠️ Неподдерживаемый формат файла: {file['name']} (ожидается .docx)")
            continue

        if not text.strip():
            not_updated.append((module, discipline))
            continue

        keywords = generate_keywords_from_text(text)
        if keywords:
            update_keywords_for_discipline(module, discipline, keywords)
            logger.info(f"✅ Обновлены ключевые слова для {discipline}")
        else:
            not_updated.append((module, discipline))

    if not_updated:
        logger.warning("\n❌ Не удалось обновить ключевые слова для:")
        for m, d in not_updated:
            logger.warning(f"  - {program_name} → {m} → {d}")
    else:
        logger.info("\n✅ Все ключевые слова успешно обновлены!")

if __name__ == "__main__":
    for program, sheet in PROGRAM_SHEETS_LIST.items():
        process_program(program, sheet)